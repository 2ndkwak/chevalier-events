"""
Avery 5011 tent card generator — pure Python, no Node required.
Avery 5011: 3-3/4" W x 2-7/8" H per card, 6 per sheet (2 cols x 3 rows)
Margins: top 1.38", left 0.75", right 0.25", bottom 0.62"

Card layout (when printed and folded on the horizontal centre):
  TOP HALF    → back face  : first name, 180° rotated, large bold
  CENTRE LINE → fold line  : subtle dashed rule
  BOTTOM HALF → front face : logo + full name + table/seat
"""
import sys, json
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree

BURGUNDY = RGBColor(0x6B, 0x1A, 0x2A)
MUTED    = RGBColor(0x7A, 0x66, 0x50)

# Card dimensions
CARD_W_IN = 3.75
CARD_H_IN = 2.875   # = 4140 twips = Avery 5011 exact cell height.
                     # Do NOT change without updating the physical Avery
                     # alignment (6-per-sheet die-cut grid).
                     #
                     # IMPORTANT: this height is enforced with hRule="exact",
                     # which Word's screen view and Print Preview both tolerate
                     # slight overflow against — but the real print/RIP path
                     # clips silently, always losing the last line in the cell
                     # (the table/seat number). Verified by hand on a real
                     # Avery 5011 sheet. Keep total stacked content (rotated
                     # name box + fold line + logo + name + table/seat line,
                     # including all cell margins and paragraph spacing)
                     # at or under ~3800 twips — i.e. leave a ~300-400 twip
                     # safety margin below the 4140 twip ceiling. Don't spec
                     # content flush to the limit, and re-verify with an
                     # actual printed sheet (not just Print Preview) after
                     # any layout change here.
HALF_H_IN = CARD_H_IN / 2   # 1.4375"

# EMU constants
EMU_PER_IN = 914400
CARD_W_EMU = int(CARD_W_IN * EMU_PER_IN)
HALF_H_EMU = int(HALF_H_IN * EMU_PER_IN)

# DXA constants (1440 per inch)
CARD_W_DXA = int(CARD_W_IN * 1440)
CARD_H_DXA = int(CARD_H_IN * 1440)


# ── XML namespaces ────────────────────────────────────────────────────────────
W   = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
WP  = 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing'
A   = 'http://schemas.openxmlformats.org/drawingml/2006/main'
WPS = 'http://schemas.microsoft.com/office/word/2010/wordprocessingShape'
MC  = 'http://schemas.openxmlformats.org/markup-compatibility/2006'


def _fix_zoom(doc):
    zoom = doc.settings.element.find(qn('w:zoom'))
    if zoom is not None and not zoom.get(qn('w:percent')):
        zoom.set(qn('w:percent'), '100')


def _nil_border(side):
    b = OxmlElement(f'w:{side}')
    b.set(qn('w:val'), 'nil')
    return b


def _set_table_no_borders(tbl):
    tblPr = tbl._tbl.tblPr
    tblBorders = OxmlElement('w:tblBorders')
    for side in ('top','left','bottom','right','insideH','insideV'):
        tblBorders.append(_nil_border(side))
    tblW = tblPr.find(qn('w:tblW'))
    if tblW is not None:
        tblW.addnext(tblBorders)
    else:
        tblPr.append(tblBorders)


def _set_cell_props(cell, width_in):
    tcPr = cell._tc.get_or_add_tcPr()
    tcW = tcPr.find(qn('w:tcW'))
    if tcW is None:
        tcW = OxmlElement('w:tcW')
        tcPr.insert(0, tcW)
    tcW.set(qn('w:w'), str(int(width_in * 1440)))
    tcW.set(qn('w:type'), 'dxa')
    # no borders — after tcW
    tcBorders = OxmlElement('w:tcBorders')
    for side in ('top','left','bottom','right'):
        tcBorders.append(_nil_border(side))
    tcW.addnext(tcBorders)
    # margins — after tcBorders
    tcMar = OxmlElement('w:tcMar')
    for side, val in (('top',120),('left',280),('bottom',60),('right',280)):
        m = OxmlElement(f'w:{side}')
        m.set(qn('w:w'), str(val))
        m.set(qn('w:type'), 'dxa')
        tcMar.append(m)
    tcBorders.addnext(tcMar)


def _set_row_height(row, inches):
    trPr = row._tr.get_or_add_trPr()
    trH = OxmlElement('w:trHeight')
    trH.set(qn('w:val'), str(int(inches * 1440)))
    trH.set(qn('w:hRule'), 'exact')
    trPr.append(trH)


def _rotated_name_paragraph(first_name):
    """
    Return a <w:p> element containing an inline drawing textbox
    with the first name rotated 180° (upside-down).
    Sized to fill the top half of the card (CARD_W x HALF_H).
    """
    # Unique ID (simple counter would be better in a real app)
    shape_id = 100

    xml = f'''<w:p xmlns:w="{W}"
               xmlns:wp="{WP}"
               xmlns:a="{A}"
               xmlns:wps="{WPS}"
               xmlns:mc="{MC}">
      <w:pPr>
        <w:jc w:val="center"/>
      </w:pPr>
      <w:r>
        <w:rPr/>
        <mc:AlternateContent>
          <mc:Choice Requires="wps">
            <w:drawing>
              <wp:inline distT="0" distB="0" distL="0" distR="0">
                <wp:extent cx="{CARD_W_EMU}" cy="{HALF_H_EMU}"/>
                <wp:effectExtent l="0" t="0" r="0" b="0"/>
                <wp:docPr id="{shape_id}" name="RotName{shape_id}"/>
                <wp:cNvGraphicFramePr>
                  <a:graphicFrameLocks xmlns:a="{A}" noChangeAspect="1"/>
                </wp:cNvGraphicFramePr>
                <a:graphic xmlns:a="{A}">
                  <a:graphicData uri="{WPS}">
                    <wps:wsp>
                      <wps:cNvSpPr txBx="1">
                        <a:spLocks noChangeArrowheads="1"/>
                      </wps:cNvSpPr>
                      <wps:spPr>
                        <a:xfrm rot="10800000">
                          <a:off x="0" y="0"/>
                          <a:ext cx="{CARD_W_EMU}" cy="{HALF_H_EMU}"/>
                        </a:xfrm>
                        <a:prstGeom prst="rect"><a:avLst/></a:prstGeom>
                        <a:noFill/>
                        <a:ln><a:noFill/></a:ln>
                      </wps:spPr>
                      <wps:txbx>
                        <w:txbxContent>
                          <w:p>
                            <w:pPr>
                              <w:jc w:val="center"/>
                              <w:spacing w:before="0" w:after="0"/>
                            </w:pPr>
                            <w:r>
                              <w:rPr>
                                <w:rFonts w:ascii="Georgia" w:hAnsi="Georgia"/>
                                <w:b/>
                                <w:color w:val="6B1A2A"/>
                                <w:sz w:val="56"/>
                                <w:szCs w:val="56"/>
                              </w:rPr>
                              <w:t>{first_name}</w:t>
                            </w:r>
                          </w:p>
                        </w:txbxContent>
                      </wps:txbx>
                      <wps:bodyPr anchor="ctr"
                                  lIns="91440" rIns="91440"
                                  tIns="45720" bIns="45720">
                        <a:noAutofit/>
                      </wps:bodyPr>
                    </wps:wsp>
                  </a:graphicData>
                </a:graphic>
              </wp:inline>
            </w:drawing>
          </mc:Choice>
          <mc:Fallback>
            <w:pict>
              <v:rect xmlns:v="urn:schemas-microsoft-com:vml"
                      style="width:{CARD_W_IN}in;height:{HALF_H_IN}in">
                <v:textbox><w:txbxContent><w:p><w:r>
                  <w:t>{first_name}</w:t>
                </w:r></w:p></w:txbxContent></v:textbox>
              </v:rect>
            </w:pict>
          </mc:Fallback>
        </mc:AlternateContent>
      </w:r>
    </w:p>'''

    return etree.fromstring(xml)


def _fill_card(cell, seat, logo_path):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

    # Remove default empty paragraph
    for p in list(cell.paragraphs):
        p._element.getparent().remove(p._element)

    first_name = seat['name'].split()[0]

    # ── BACK FACE: rotated first name (top half) ──────────────────────────
    rot_p = _rotated_name_paragraph(first_name)
    cell._tc.append(rot_p)

    # ── FOLD LINE: subtle rule ─────────────────────────────────────────────
    p_fold = cell.add_paragraph()
    p_fold.paragraph_format.space_before = Pt(0)
    p_fold.paragraph_format.space_after  = Pt(0)
    pPr = p_fold._p.get_or_add_pPr()
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:before'), '0')
    spacing.set(qn('w:after'),  '0')
    spacing.set(qn('w:line'),   '80')   # 4pt — exact, minimal; the dashed
    spacing.set(qn('w:lineRule'), 'exact')  # border draws regardless of line height
    pPr.append(spacing)
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'dashed')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'BBAA99')
    pBdr.append(bottom)
    pPr.insert(0, pBdr)
    p_fold.add_run(' ').font.size = Pt(1)

    # ── FRONT FACE: logo + full name + table/seat (bottom half) ───────────
    p_logo = cell.add_paragraph()
    p_logo.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_logo.paragraph_format.space_before = Pt(2)
    p_logo.paragraph_format.space_after  = Pt(1.5)
    p_logo.add_run().add_picture(logo_path, width=Inches(0.38))

    if seat.get('has_allergy'):
        # Right-aligned tab stop at the far edge of the usable cell width
        # (card width minus the left/right cell margins set in
        # _set_cell_props), so the dot sits in the top-right corner of the
        # front face -- big and out of the way, easy for a server to spot
        # at a glance rather than blending into the name line below.
        usable_width_in = CARD_W_IN - 2 * (280 / 1440)
        tab_stops = p_logo.paragraph_format.tab_stops
        tab_stops.add_tab_stop(Inches(usable_width_in - 0.05), WD_TAB_ALIGNMENT.RIGHT)
        p_logo.add_run('\t')
        r_dot = p_logo.add_run('\u25CF')   # red dot -- active allergy flag for this event
        r_dot.font.name = 'Georgia'; r_dot.font.size = Pt(22); r_dot.font.bold = True
        r_dot.font.color.rgb = RGBColor(0xC4, 0x1E, 0x1E)

    p_name = cell.add_paragraph()
    p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_name.paragraph_format.space_before = Pt(0)
    p_name.paragraph_format.space_after  = Pt(1)
    r = p_name.add_run(seat['name'])
    r.font.name = 'Georgia'; r.font.size = Pt(15)
    r.font.color.rgb = BURGUNDY; r.font.bold = False

    p_loc = cell.add_paragraph()
    p_loc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_loc.paragraph_format.space_before = Pt(0)
    p_loc.paragraph_format.space_after  = Pt(0)
    r2 = p_loc.add_run(f"{seat['table_label']}  ·  Seat {seat['seat_num']}")
    r2.font.name = 'Arial'; r2.font.size = Pt(7)
    r2.font.color.rgb = MUTED; r2.font.all_caps = True


def _section_break_para(doc):
    """Near-invisible paragraph carrying a page-section break."""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    # spacing BEFORE sectPr in pPr
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:before'), '0')
    spacing.set(qn('w:after'),  '0')
    spacing.set(qn('w:line'),   '20')
    spacing.set(qn('w:lineRule'), 'exact')
    pPr.append(spacing)
    # sectPr
    sectPr = OxmlElement('w:sectPr')
    pgSz = OxmlElement('w:pgSz')
    pgSz.set(qn('w:w'), str(int(8.5 * 1440)))
    pgSz.set(qn('w:h'), str(int(11  * 1440)))
    sectPr.append(pgSz)
    pgMar = OxmlElement('w:pgMar')
    pgMar.set(qn('w:top'),    str(int(1.38 * 1440)))
    pgMar.set(qn('w:bottom'), str(int(0.62 * 1440)))
    pgMar.set(qn('w:left'),   str(int(0.75 * 1440)))
    pgMar.set(qn('w:right'),  str(int(0.25 * 1440)))
    pgMar.set(qn('w:header'), '720')
    pgMar.set(qn('w:footer'), '720')
    pgMar.set(qn('w:gutter'), '0')
    sectPr.append(pgMar)
    pPr.append(sectPr)
    r = p.add_run(' ')
    r.font.size = Pt(1)


def generate(seats_data, logo_path, out_path):
    sorted_seats = sorted(
        seats_data,
        key=lambda s: (s.get('table_num', 0), s['seat_num'])
    )
    while len(sorted_seats) % 6 != 0:
        sorted_seats.append(None)

    doc = Document()
    _fix_zoom(doc)

    for p in list(doc.paragraphs):
        p._element.getparent().remove(p._element)

    sec = doc.sections[0]
    sec.page_width    = Inches(8.5)
    sec.page_height   = Inches(11)
    sec.top_margin    = Inches(1.38)
    sec.bottom_margin = Inches(0.62)
    sec.left_margin   = Inches(0.75)
    sec.right_margin  = Inches(0.25)

    num_sheets = len(sorted_seats) // 6
    for sheet_idx in range(num_sheets):
        chunk = sorted_seats[sheet_idx * 6 : sheet_idx * 6 + 6]

        if sheet_idx > 0:
            _section_break_para(doc)

        tbl = doc.add_table(rows=3, cols=2)
        _set_table_no_borders(tbl)

        for r in range(3):
            _set_row_height(tbl.rows[r], CARD_H_IN)
            for c in range(2):
                cell = tbl.cell(r, c)
                _set_cell_props(cell, CARD_W_IN)
                seat = chunk[r * 2 + c]
                if seat:
                    _fill_card(cell, seat, logo_path)

    doc.save(out_path)


if __name__ == '__main__':
    generate(json.loads(sys.argv[1]), sys.argv[3], sys.argv[2])
    print("OK")
